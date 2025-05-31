import os
from langchain_community.document_loaders import TextLoader, PyPDFLoader
from langchain_core.documents import Document
from docx import Document as DocxDocument
import pandas as pd

class DataLoader:
    def __init__(self, folder_path):
        self.folder_path = folder_path
    
    def load_markdown(self, path):
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()
        return [Document(page_content=content, metadata={"source": path})]

    def text_loader(self, path):
        try:
            # Try the LangChain TextLoader first
            loader = TextLoader(file_path=path)
            return loader.load()
        except Exception as e:
            print(f"TextLoader error: {str(e)}, trying direct reading")
            try:
                # Fallback: read file directly
                with open(path, 'r', encoding='utf-8') as f:
                    content = f.read()
                return [Document(page_content=content, metadata={"source": path})]
            except UnicodeDecodeError:
                # Try another encoding if UTF-8 fails
                with open(path, 'r', encoding='latin-1') as f:
                    content = f.read()
                return [Document(page_content=content, metadata={"source": path})]

    def pdf_loader(self, path):
        loader = PyPDFLoader(file_path=path)
        return loader.load()

    def xlsx_csv_loader(self, path):
        try:
            if path.lower().endswith('.csv'):
                df = pd.read_csv(path)
            else:  # For xlsx files
                df = pd.read_excel(path)
            content = df.to_string()
            return [Document(page_content=content, metadata={"source": path})]
        except Exception as e:
            print(f"Error loading {path}: {str(e)}")
            return [Document(page_content=f"Error loading file: {str(e)}", metadata={"source": path})]

    def docx_loader(self, path):
        doc = DocxDocument(path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return [Document(page_content=text, metadata={"source": path})]

    def load_file(self, file_paths):
        data = []
        for path in file_paths:
            if not os.path.exists(path):
                print(f"Warning: File {path} does not exist")
                continue
                
            ext = os.path.splitext(path)[-1].lower()
            try:
                if ext == ".txt":
                    data.extend(self.text_loader(path))
                elif ext == ".pdf":
                    data.extend(self.pdf_loader(path))
                elif ext == ".docx":
                    data.extend(self.docx_loader(path))
                elif ext in [".csv", ".xlsx"]:
                    data.extend(self.xlsx_csv_loader(path))
                elif ext == ".md":
                    data.extend(self.load_markdown(path))
                else:
                    print(f"Unsupported file type: {ext}")
            except Exception as e:
                print(f"Error processing {path}: {str(e)}")
        return data

    def get_docs(self):
        docs = []
        for file_name in os.listdir(self.folder_path):
            full_path = os.path.join(self.folder_path, file_name)
            if os.path.isfile(full_path) and not file_name.startswith("~$"): 
                docs.append(full_path)
        return docs

    def load_all_files(self):
        return self.load_file(self.get_docs())